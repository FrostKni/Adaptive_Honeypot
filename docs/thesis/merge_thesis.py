#!/usr/bin/env python3
"""
Merge all thesis parts into a single complete document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def merge_documents(part_files, output_file):
    """Merge multiple docx files into one"""
    # Create a new document for the merged result
    merged_doc = Document()
    
    # Copy styles from first document
    first_doc = Document(part_files[0])
    
    # Process each part
    for i, part_file in enumerate(part_files):
        print(f"Processing {part_file}...")
        doc = Document(part_file)
        
        # Copy all content from this document
        for element in doc.element.body:
            merged_doc.element.body.append(element)
    
    # Save the merged document
    merged_doc.save(output_file)
    print(f"\nMerged document saved to: {output_file}")
    
    # Get file size
    size = os.path.getsize(output_file)
    print(f"File size: {size / 1024:.1f} KB")
    
    return output_file

def count_pages(docx_file):
    """Estimate page count based on content"""
    doc = Document(docx_file)
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    
    # Rough estimate: ~30 paragraphs per page
    estimated_pages = paragraphs // 30 + tables
    
    print(f"Estimated content:")
    print(f"  - Paragraphs: {paragraphs}")
    print(f"  - Tables: {tables}")
    print(f"  - Estimated pages: ~{estimated_pages}")
    
    return estimated_pages

if __name__ == "__main__":
    thesis_dir = "/home/kali/Music/Adaptive_Honeypot/docs/thesis"
    
    part_files = [
        f"{thesis_dir}/Adaptive_Honeypot_Thesis_Part1.docx",
        f"{thesis_dir}/Adaptive_Honeypot_Thesis_Part2.docx",
        f"{thesis_dir}/Adaptive_Honeypot_Thesis_Part3.docx",
    ]
    
    output_file = f"{thesis_dir}/Adaptive_Honeypot_Thesis_Complete.docx"
    
    # Check all files exist
    for f in part_files:
        if os.path.exists(f):
            size = os.path.getsize(f)
            print(f"Found: {f} ({size / 1024:.1f} KB)")
        else:
            print(f"Missing: {f}")
            exit(1)
    
    print("\n" + "="*50)
    print("Merging thesis parts...")
    print("="*50 + "\n")
    
    merge_documents(part_files, output_file)
    
    print("\n" + "="*50)
    print("Content summary:")
    print("="*50)
    count_pages(output_file)
    
    print("\n" + "="*50)
    print("THESIS GENERATION COMPLETE!")
    print("="*50)
    print(f"\nOutput: {output_file}")